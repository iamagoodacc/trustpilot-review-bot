import json
import os
import time
from dotenv import load_dotenv
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_openai import AzureOpenAIEmbeddings, AzureChatOpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
from textblob import TextBlob
from langchain.docstore.document import Document as LangchainDocument

# Load environment variables from .env file
load_dotenv()

# Azure OpenAI Client for LLM
llm = AzureChatOpenAI(
    openai_api_key=os.environ["AZURE_OPENAI_API_KEY"],
    openai_api_version=os.environ["AZURE_OPENAI_CHAT_API_VERSION"],
    azure_deployment=os.environ["AZURE_OPENAI_CHAT_DEPLOYMENT_NAME"],
    azure_endpoint=os.environ["AZURE_OPENAI_ENDPOINT"],
    temperature=0.1,
)

# Azure OpenAI Embedding Client
embeddings = AzureOpenAIEmbeddings(
    azure_deployment=os.environ["AZURE_OPENAI_TEXT_EMBEDDING_DEPLOYMENT_NAME"],
    openai_api_version=os.environ["AZURE_OPENAI_TEXT_EMBEDDING_API_VERSION"],
)

# Text Splitter
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)

# Function to load and chunk text data from the file
def load_text(file_path):
    json_filename = os.path.join(file_path)
    
    if not os.path.exists(json_filename):
        print(f"File {json_filename} does not exist.")
        return None
    
    with open(json_filename, 'r') as file:
        reviews_data = json.load(file)
    
    review_text = ""
    for review in reviews_data:
        review_text += review["title"] + " " + review["body"] + "\n"
    
    return review_text

# Function to analyze sentiment
def analyze_sentiment(reviews):
    positive_reviews = 0
    negative_reviews = 0
    
    for review in reviews:
        text = review["title"] + " " + review["body"]
        sentiment = TextBlob(text).sentiment.polarity
        if sentiment > 0:
            positive_reviews += 1
        elif sentiment < 0:
            negative_reviews += 1
    
    total_reviews = len(reviews)
    positive_percentage = (positive_reviews / total_reviews) * 100
    negative_percentage = (negative_reviews / total_reviews) * 100
    
    return positive_percentage, negative_percentage

# Function to generate detailed responses using LLM and retrieved documents
def generate_responses(questions, document_search, llm_chain):
    responses = {}
    for question in questions:
        try:
            docs = document_search.similarity_search(question)
            answer = llm_chain.run(input_documents=docs, question=question)
            responses[question] = answer
        except Exception as e:
            print(f"Error generating response for question '{question}': {e}")
            responses[question] = "Error generating response."
    return responses

# Function to summarize responses for each question
def summarize_responses(responses, llm_chain):
    summarized_responses = {}
    for question, response_list in responses.items():
        combined_response = " ".join(response_list)
        doc = LangchainDocument(page_content=combined_response)
        summarized_responses[question] = llm_chain.run(input_documents=[doc], question=f"Summarize the key points regarding {question}.")
    return summarized_responses

# Function to create a Word document from the responses
def create_word_report(summarized_responses, positive_percentage, negative_percentage, output_path='analytical_report.docx'):
    document = Document()
    document.add_heading('Analytical Report', 0)

    document.add_heading('Executive Summary', level=1)
    document.add_paragraph(
        f"This comprehensive report evaluates 2,420 customer reviews of Nationwide, focusing on key areas of customer experience and satisfaction. "
        f"The analysis provides insights into several aspects of Nationwide's services, highlighting areas of strength and opportunities for improvement."
    )
    document.add_heading('Sentiment Analysis', level=2)
    document.add_paragraph(f"Positive Reviews: {positive_percentage:.2f}%")
    document.add_paragraph(f"Negative Reviews: {negative_percentage:.2f}%")

    for question, summary in summarized_responses.items():
        document.add_heading(question, level=1)
        document.add_paragraph(summary)
    
    document.save(output_path)
    print(f"Report saved as '{output_path}'.")

if __name__ == "__main__":
    # Load reviews from the existing JSON file
    review_text = load_text('reviews/nationwide.json')
    if review_text is None:
        exit(1)
    
    # Load reviews data
    with open('reviews/nationwide.json', 'r') as file:
        reviews_data = json.load(file)
    
    # Analyze sentiment
    positive_percentage, negative_percentage = analyze_sentiment(reviews_data)
    
    # Chunk the reviews text
    texts = text_splitter.split_text(review_text)

    # Batch processing to avoid rate limits
    batch_size = 100
    aggregated_responses = {question: [] for question in [
        "Customer satisfaction with branch staff and service quality",
        "Customer loyalty and retention rates for branch customers",
        "Customer feedback on the Fairer Share Payment and other member benefits",
        "Customer engagement with branch events and charitable activities",
        "Customer complaints and resolutions for branch-related issues",
        "Customer preferences and expectations for branch versus digital channels",
        "Customer awareness and perception of the mutual ownership model and its advantages"
    ]}

    questions = list(aggregated_responses.keys())

    for i in range(0, len(texts), batch_size):
        batch_texts = texts[i:i+batch_size]
        print(f"Processing batch {i // batch_size + 1}/{(len(texts) // batch_size) + 1}")

        try:
            # Create FAISS vector store for the current batch
            document_search = FAISS.from_texts(batch_texts, embeddings)
            llm_chain = load_qa_chain(llm)
            
            # Generate responses for the current batch
            responses = generate_responses(questions, document_search, llm_chain)
            for question, response in responses.items():
                aggregated_responses[question].append(response)
        
        except Exception as e:
            print(f"Error in batch {i // batch_size + 1}: {e}")
            if "429" in str(e):
                print("Rate limit hit, sleeping for 60 seconds")
                time.sleep(60)
            else:
                exit(1)
    
    # Summarize responses
    summarized_responses = summarize_responses(aggregated_responses, llm_chain)
    
    # Create the Word report
    create_word_report(summarized_responses, positive_percentage, negative_percentage)
    print("Analysis complete. Results saved to 'analytical_report.docx'.")



